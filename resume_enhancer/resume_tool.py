# Required Libraries
import os
import re
import docx
import fitz  # PyMuPDF
import streamlit as st
from docx import Document
from docx.shared import Pt

# Extract text from DOCX
def extract_text_docx(file):
    doc = docx.Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Extract text from PDF
def extract_text_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

# Basic Rule-Based Categorizer
def categorize_text(text):
    categories = {
        "Contact Information": "",
        "Education": "",
        "Experience": "",
        "Projects": "",
        "Skills": "",
        "Achievements": "",
        "Certifications": "",
        "Others": ""
    }

    lines = text.split('\n')
    current_category = "Others"
    
    for line in lines:
        line = line.strip()
        if not line:
            continue

        lower_line = line.lower()

        if any(x in lower_line for x in ["education", "qualifications"]):
            current_category = "Education"
        elif any(x in lower_line for x in ["experience", "employment", "work"]):
            current_category = "Experience"
        elif any(x in lower_line for x in ["project"]):
            current_category = "Projects"
        elif any(x in lower_line for x in ["skill"]):
            current_category = "Skills"
        elif any(x in lower_line for x in ["achievement", "award", "honor"]):
            current_category = "Achievements"
        elif any(x in lower_line for x in ["certification", "course"]):
            current_category = "Certifications"
        elif any(x in lower_line for x in ["phone", "email", "linkedin", "github"]):
            current_category = "Contact Information"

        categories[current_category] += line + '\n'

    return categories

# Create Word File from Categories
def create_word_file(categories):
    doc = Document()
    for category, content in categories.items():
        if content.strip():
            doc.add_heading(category, level=1)
            for line in content.strip().split('\n'):
                doc.add_paragraph(line, style='List Bullet')
    output_path = "categorized_resume.docx"
    doc.save(output_path)
    return output_path

# Streamlit UI
st.title("📄 Resume Categorizer Tool")

uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file is not None:
    if uploaded_file.name.endswith(".pdf"):
        raw_text = extract_text_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        raw_text = extract_text_docx(uploaded_file)
    else:
        st.error("Unsupported file format.")
        st.stop()

    st.success("✅ Resume uploaded and text extracted.")

    categorized = categorize_text(raw_text)
    output_file_path = create_word_file(categorized)

    with open(output_file_path, "rb") as file:
        st.download_button(
            label="📥 Download Categorized Resume",
            data=file,
            file_name="categorized_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )